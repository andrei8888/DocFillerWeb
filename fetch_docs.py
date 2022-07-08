import os
import re
import shutil
from datetime import date

from docx import Document
from docx.shared import Inches

import informations

DOCS_PATH = "static/docs/"
DOCS_TEMPLATES_PATH = DOCS_PATH + "doc_templates/"
DOCS_UPLOAD_PATH = "static/uploads/docs_completed/"
FINAL_ZIP_NAME = "Documente_completate"
FINAL_ZIP_NAME_EXT = FINAL_ZIP_NAME + ".zip"


def fetch_titles_and_number():
    list_docs = [os.path.splitext(filename)[0] for filename in os.listdir(DOCS_PATH) if
                 os.path.isfile(os.path.join(DOCS_PATH, filename))]
    # print(list_docs)
    return list_docs, len(list_docs)


def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text


def docx_replace_text_with_image(doc_obj, regex):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub("", inline[i].text)
                    inline[i].text = text
                    inline[i].add_picture(f'{"static/uploads/"+informations.signature_file}', width=Inches(1), height=Inches(0.5))


def get_current_date():
    return date.today().strftime("%d.%m.%Y")


def replace_with_infos(doc_obj, infos):
    id_char = "@"
    infos["domiciliu"] = infos["domiciliu"].replace("\n", " ")
    for tipInfo in infos.keys():
        docx_replace_regex(doc_obj, re.compile(id_char+tipInfo), infos[tipInfo])
    docx_replace_regex(doc_obj, re.compile(id_char+"dataCurenta"), get_current_date())
    docx_replace_text_with_image(doc_obj, re.compile(id_char + "semnatura"))


def manage_doc_and_save(doc_name, infos):
    curr_doc = Document(DOCS_PATH + "/doc_templates/" + doc_name)
    replace_with_infos(curr_doc, infos)
    curr_doc.save(DOCS_UPLOAD_PATH + doc_name)


def make_zip():
    shutil.make_archive(FINAL_ZIP_NAME, 'zip', DOCS_UPLOAD_PATH)
    shutil.move(FINAL_ZIP_NAME_EXT, DOCS_UPLOAD_PATH + FINAL_ZIP_NAME_EXT)


def clean_docs_completed():
    for f in os.listdir(DOCS_UPLOAD_PATH):
        os.remove(os.path.join(DOCS_UPLOAD_PATH, f))


def manage_documents(docs):
    for doc_title in docs.values():
        manage_doc_and_save(doc_title + ".docx", informations.person_informations)

    make_zip()
